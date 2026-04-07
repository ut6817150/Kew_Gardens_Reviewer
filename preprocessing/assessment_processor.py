"""
assessment_processor.py

Convert assessment files from `.docx` or `.html` into a structured JSON
dictionary with headings, paragraphs, lists, tables, and comments.

Run:
  Batch mode:
    python3.12 assessment_processor.py
    Reads files from `converted/` and saves JSON files to `json_converted/`.

  Single file:
    python3.12 assessment_processor.py "converted/My Assessment.docx"
    Prints the parsed JSON to stdout.
"""

# Imports
from __future__ import annotations
import json
import sys
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple, Union
from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from bs4 import BeautifulSoup


# Default folders
DEFAULT_INPUT_FOLDER = "converted"
DEFAULT_OUTPUT_FOLDER = "json_converted"


# Data structure
@dataclass
class HeadingNode:
    """A heading node in the skeleton tree."""
    title: Optional[str]
    level: int
    path: List[str]
    blocks: List[Dict[str, Any]]
    children: List["HeadingNode"]

    def to_dict(self) -> Dict[str, Any]:
        return {
            "title": self.title,
            "level": self.level,
            "path": list(self.path),
            "blocks": list(self.blocks),
            "children": [c.to_dict() for c in self.children],
        }


class AssessmentParser:
    """
    Parser for assessment documents.

    It builds the heading structure, attaches content blocks, and extracts
    comments with anchor context.
    """

    # WordprocessingML namespace used by DOCX comment and anchor XML.
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS = {"w": W_NS}

    def parse_file(self, path: str) -> Dict[str, Any]:
        """Parse one `.docx` or `.html` file into the output JSON schema."""
        p = Path(path)
        ext = p.suffix.lower()

        if ext == ".docx":
            doc = Document(path)
            doc_title = p.stem
            root = self._build_heading_skeleton(doc, doc_title=doc_title)
            self._attach_blocks(doc, root)

            out = root.to_dict()
            out["comments"] = self._extract_comments_with_anchors(path)
            return out

        if ext in (".html", ".htm"):
            return self._parse_html(path)

        raise ValueError(f"Unsupported file type: {ext}")
    # DOCX traversal
    def _iter_block_items_in_order(self, doc: _Document) -> Iterable[Union[Paragraph, Table]]:
        """Yield top-level paragraphs and tables in the order they appear."""
        for child in doc.element.body.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, doc)
            elif child.tag.endswith("}tbl"):
                yield Table(child, doc)

    def _paragraph_style_name(self, p: Paragraph) -> str:
        """Return the paragraph style name, or an empty string."""
        try:
            return (p.style.name or "").strip()
        except Exception:
            return ""

    def _heading_level(self, p: Paragraph) -> Optional[int]:
        """Return 1/2 if style is Heading 1/Heading 2 else None."""
        name = self._paragraph_style_name(p)
        if name == "Heading 1" or name.replace(" ", "").lower() == "heading1":
            return 1
        if name == "Heading 2" or name.replace(" ", "").lower() == "heading2":
            return 2
        return None

    def _is_list_item(self, p: Paragraph) -> bool:
        """Detect Word list items via <w:numPr>."""
        ppr = p._p.pPr
        return ppr is not None and ppr.numPr is not None

    def _list_signature(self, p: Paragraph) -> str:
        """Group consecutive list items by numId/ilvl where possible."""
        ppr = p._p.pPr
        style = self._paragraph_style_name(p)

        if ppr is not None and ppr.numPr is not None:
            numPr = ppr.numPr
            numId = numPr.numId.val if numPr.numId is not None else None
            ilvl = numPr.ilvl.val if numPr.ilvl is not None else None
            return f"numId={numId}|ilvl={ilvl}|style={style}"

        return f"style={style}"

    def _table_to_rows(self, tbl: Table) -> List[List[str]]:
        """Extract table into 2D list of cell text."""
        rows: List[List[str]] = []
        for row in tbl.rows:
            out_row: List[str] = []
            for cell in row.cells:
                text = "\n".join(
                    [p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()]
                )
                out_row.append(text)
            rows.append(out_row)
        return rows

    def _table_to_rows_rich(self, tbl: Table) -> List[List[str]]:
        """Extract table into 2D list of cell text preserving inline rich-text markup."""
        rows: List[List[str]] = []
        for row in tbl.rows:
            out_row: List[str] = []
            for cell in row.cells:
                text = "\n".join(
                    [
                        rich_text
                        for p in cell.paragraphs
                        if (rich_text := self._paragraph_text_with_scripts(p)).strip()
                    ]
                )
                out_row.append(text)
            rows.append(out_row)
        return rows

    def _empty_style_bucket(self) -> Dict[str, List[str]]:
        """Return an empty container for collected style snippets."""
        return {"bold": [], "italic": []}

    def _run_vertical_align(self, run) -> Optional[str]:
        """
        Return 'superscript', 'subscript', or None for a run.

        Prefer python-docx properties when available, then fall back to the raw
        WordprocessingML vertAlign value.
        """
        try:
            if run.font.superscript is True:
                return "superscript"
            if run.font.subscript is True:
                return "subscript"
        except Exception:
            pass

        try:
            vert = run._r.find(".//w:vertAlign", self.NS)
            if vert is not None:
                val = vert.get(f"{{{self.W_NS}}}val")
                if val == "superscript":
                    return "superscript"
                if val == "subscript":
                    return "subscript"
        except Exception:
            pass

        return None

    def _effective_run_bool(self, p: Paragraph, run, prop_name: str) -> bool:
        """Resolve a run boolean format, inheriting from run and paragraph styles."""
        def tri_to_bool(x):
            return None if x is None else bool(x)

        direct = tri_to_bool(getattr(run, prop_name))
        if direct is not None:
            return direct

        try:
            if run.style and run.style.font:
                v = tri_to_bool(getattr(run.style.font, prop_name))
                if v is not None:
                    return v
        except Exception:
            pass

        try:
            if p.style and p.style.font:
                v = tri_to_bool(getattr(p.style.font, prop_name))
                if v is not None:
                    return v
        except Exception:
            pass

        return False

    def _wrap_rich_text(self, text: str, *, bold: bool = False, italic: bool = False, vertical: Optional[str] = None) -> str:
        """Wrap a text chunk with inline rich-text tags in a stable order."""
        rendered = text

        if vertical == "superscript":
            rendered = f"<sup>{rendered}</sup>"
        elif vertical == "subscript":
            rendered = f"<sub>{rendered}</sub>"

        if italic:
            rendered = f"<i>{rendered}</i>"
        if bold:
            rendered = f"<b>{rendered}</b>"

        return rendered

    def _paragraph_text_with_scripts(self, p: Paragraph) -> str:
        """
        Return paragraph text while preserving bold/italic and superscript/
        subscript using inline HTML-like markup.
        """
        parts: List[str] = []
        cur_fmt: Optional[Tuple[bool, bool, Optional[str]]] = None
        buf: List[str] = []

        def flush() -> None:
            nonlocal buf, cur_fmt
            if cur_fmt is None or not buf:
                return
            text = "".join(buf)
            parts.append(
                self._wrap_rich_text(
                    text,
                    bold=cur_fmt[0],
                    italic=cur_fmt[1],
                    vertical=cur_fmt[2],
                )
            )
            buf = []

        for run in getattr(p, "runs", []):
            if not run.text:
                continue

            fmt = (
                self._effective_run_bool(p, run, "bold"),
                self._effective_run_bool(p, run, "italic"),
                self._run_vertical_align(run),
            )
            if cur_fmt is None:
                cur_fmt = fmt
            if fmt != cur_fmt:
                flush()
                cur_fmt = fmt

            buf.append(run.text)

        flush()

        return "".join(parts).strip()

    def _xml_run_bool(self, run_el: ET.Element, prop_name: str) -> bool:
        """Read a direct boolean run property from raw WordprocessingML."""
        prop = run_el.find(f"./w:rPr/w:{prop_name}", self.NS)
        if prop is None:
            return False

        val = prop.get(f"{{{self.W_NS}}}val")
        if val is None:
            return True
        return val.lower() not in {"0", "false", "off"}

    def _xml_run_text_with_scripts(self, run_el: ET.Element) -> str:
        """Render a raw WordprocessingML <w:r> element with inline rich-text markup."""
        texts = [t.text for t in run_el.findall(".//w:t", self.NS) if t.text]
        if not texts:
            return ""

        text = "".join(texts)
        vert = run_el.find(".//w:vertAlign", self.NS)
        vert_val = vert.get(f"{{{self.W_NS}}}val") if vert is not None else None
        if vert_val not in {"superscript", "subscript"}:
            vert_val = None

        return self._wrap_rich_text(
            text,
            bold=self._xml_run_bool(run_el, "b"),
            italic=self._xml_run_bool(run_el, "i"),
            vertical=vert_val,
        )

    def _merge_style_bucket(self, target: Dict[str, List[str]], src: Dict[str, List[str]]) -> None:
        # Preserve order and avoid duplicates.
        for k, vals in src.items():
            if k not in target:
                target[k] = []
            for v in vals:
                v = (v or "").strip()
                if v and v not in target[k]:
                    target[k].append(v)

    def _extract_styles_from_docx_paragraph(self, p: Paragraph) -> Dict[str, List[str]]:
        """
        Extract formatted snippets from a DOCX paragraph using effective formatting:
        - run.bold/italic/underline when explicitly set
        - otherwise inherit from run.style.font then paragraph style

        Also merges contiguous runs with the same formatting so multi-run phrases
        become ONE string.
        """
        out = self._empty_style_bucket()

        def flush(flags, buf):
            txt = "".join(buf).replace("\xa0", " ").strip()
            if not txt:
                return
            bold, italic = flags
            if bold:
                out["bold"].append(txt)
            if italic:
                out["italic"].append(txt)

        cur_flags = None
        buf: List[str] = []

        for run in getattr(p, "runs", []):
            # Keep per-run spacing so punctuation and words join correctly.
            t = run.text
            if not t:
                continue

            flags = (
                self._effective_run_bool(p, run, "bold"),
                self._effective_run_bool(p, run, "italic"),
            )

            if cur_flags is None:
                cur_flags = flags

            if flags != cur_flags:
                flush(cur_flags, buf)
                buf = []
                cur_flags = flags

            buf.append(t)

        if cur_flags is not None:
            flush(cur_flags, buf)

        # Hyperlinks are collected separately from the raw XML.
        try:
            for hl in p._p.findall(".//w:hyperlink", self.NS):
                texts = [t.text for t in hl.findall(".//w:t", self.NS) if t.text]
                chunk = "".join(texts).replace("\xa0", " ").strip()
                if chunk:
                    out["hyperlink"].append(chunk)
        except Exception:
            pass

        deduped = self._empty_style_bucket()
        self._merge_style_bucket(deduped, out)
        return deduped


    def _extract_styles_from_html_element(self, el, bold_classes: set[str] | None = None, italic_classes: set[str] | None = None) -> Dict[str, List[str]]:
        """
        Extract formatted snippets from an HTML element using tags + inline style.
        Looks for: strong/b, em/i, u, a, code/pre, and inline styles like font-weight, font-style.
        """
        bold_classes = set(bold_classes or [])
        italic_classes = set(italic_classes or [])

        out = self._empty_style_bucket()

        def clean(s: str) -> str:
            return " ".join((s or "").split())

        # Tag-based
        for tag in el.find_all(["strong", "b"]):
            t = clean(tag.get_text(" ", strip=True))
            if t:
                out["bold"].append(t)

        for tag in el.find_all(["em", "i"]):
            t = clean(tag.get_text(" ", strip=True))
            if t:
                out["italic"].append(t)

        # Class-based (common in SIS HTML: <span class="dataLabel">...</span>)
        BOLD_CLASSES = {"dataLabel"}   
        ITALIC_CLASSES = set()         

        for tag in el.find_all(True):
            classes = set(tag.get("class") or [])
            if not classes:
                continue

            t = clean(tag.get_text(" ", strip=True))
            if not t:
                continue

            if classes & BOLD_CLASSES:
                out["bold"].append(t)
            if classes & ITALIC_CLASSES:
                out["italic"].append(t)

        # Inline style-based (covers spans like <span style="font-weight:700">)
        for tag in el.find_all(True):
            style = (tag.get("style") or "").lower()
            if not style:
                continue
            t = clean(tag.get_text(" ", strip=True))
            if not t:
                continue

            if "font-weight" in style and ("bold" in style or "700" in style or "800" in style or "900" in style):
                out["bold"].append(t)
            if "font-style" in style and "italic" in style:
                out["italic"].append(t)

        deduped = self._empty_style_bucket()
        self._merge_style_bucket(deduped, out)
        return deduped

    # Heading builder
    def _build_heading_skeleton(self, doc: _Document, doc_title: str) -> HeadingNode:
        """
        Build Heading 1/Heading 2 skeleton in document order.
        """
        root = HeadingNode(title=doc_title, level=0, path=[], blocks=[], children=[])
        current_h1: Optional[HeadingNode] = None

        def add_heading(title: str, lvl: int) -> None:
            nonlocal current_h1
            title = title.strip()
            if not title:
                return

            if lvl == 1:
                node = HeadingNode(title=title, level=1, path=[title], blocks=[], children=[])
                root.children.append(node)
                current_h1 = node
                return

            if lvl == 2:
                parent = current_h1
                node_path = [title] if parent is None else parent.path + [title]
                node = HeadingNode(title=title, level=2, path=node_path, blocks=[], children=[])
                if parent is None:
                    root.children.append(node)  # orphan H2
                else:
                    parent.children.append(node)

        for item in self._iter_block_items_in_order(doc):
            if isinstance(item, Paragraph):
                lvl = self._heading_level(item)
                if lvl in (1, 2):
                    add_heading((item.text or "").strip(), lvl)

        return root

    # Block attachment
    def _attach_blocks(self, doc: _Document, root: HeadingNode) -> None:
        """
        Walk the doc in order and attach paragraph/list/table blocks to the
        current heading (deepest of H2 then H1; else root).
        """
        current_h1: Optional[HeadingNode] = None
        current_h2: Optional[HeadingNode] = None

        def container() -> HeadingNode:
            if current_h2 is not None:
                return current_h2
            if current_h1 is not None:
                return current_h1
            return root
        
        def style_bucket_for(node: HeadingNode) -> Dict[str, List[str]]:
            if not hasattr(node, "_style_bucket"):
                setattr(node, "_style_bucket", self._empty_style_bucket())
            return getattr(node, "_style_bucket")

        def set_heading(title: str, lvl: int) -> None:
            nonlocal current_h1, current_h2
            title = title.strip()
            if not title:
                return

            if lvl == 1:
                for n in root.children:
                    if n.level == 1 and n.title == title:
                        current_h1, current_h2 = n, None
                        return
                # Fallback: create the heading if it was not in the skeleton.
                n = HeadingNode(title=title, level=1, path=[title], blocks=[], children=[])
                root.children.append(n)
                current_h1, current_h2 = n, None
                return

            if lvl == 2:
                parent = current_h1
                if parent is None:
                    # Allow Heading 2 items that appear before any Heading 1.
                    for n in root.children:
                        if n.level == 2 and n.title == title:
                            current_h2 = n
                            return
                    n = HeadingNode(title=title, level=2, path=[title], blocks=[], children=[])
                    root.children.append(n)
                    current_h2 = n
                    return

                for n in parent.children:
                    if n.level == 2 and n.title == title:
                        current_h2 = n
                        return
                n = HeadingNode(title=title, level=2, path=parent.path + [title], blocks=[], children=[])
                parent.children.append(n)
                current_h2 = n

        pending_list: Optional[Dict[str, Any]] = None
        pending_sig: Optional[str] = None

        def flush_list() -> None:
            nonlocal pending_list, pending_sig
            if pending_list is not None:
                container().blocks.append(pending_list)
                pending_list = None
                pending_sig = None

        for item in self._iter_block_items_in_order(doc):
            if isinstance(item, Paragraph):
                lvl = self._heading_level(item)
                if lvl in (1, 2):
                    flush_list()
                    set_heading(item.text, lvl)
                    continue

                text = (item.text or "").strip()
                rich_text = self._paragraph_text_with_scripts(item)
                if not text:
                    continue

                self._merge_style_bucket(style_bucket_for(container()), self._extract_styles_from_docx_paragraph(item))


                if self._is_list_item(item):
                    sig = self._list_signature(item)
                    if pending_list is None or sig != pending_sig:
                        flush_list()
                        pending_list = {"type": "list", "signature": sig, "items": []}
                        pending_sig = sig
                    pending_list["items"].append(
                        {
                            "text": text,
                            "text_rich": rich_text,
                            "style": self._paragraph_style_name(item),
                        }
                    )
                else:
                    flush_list()
                    container().blocks.append({"type": "paragraph", "text": text, "text_rich": rich_text})

            elif isinstance(item, Table):
                flush_list()
                container().blocks.append(
                    {
                        "type": "table",
                        "rows": self._table_to_rows(item),
                        "rows_rich": self._table_to_rows_rich(item),
                    }
                )

        flush_list()

        def add_style_block_if_any(node: HeadingNode) -> None:
            bucket = getattr(node, "_style_bucket", None)
            if bucket and any(bucket[k] for k in bucket):
                node.blocks.append({"type": "style", "data": bucket})
            for ch in node.children:
                add_style_block_if_any(ch)

        add_style_block_if_any(root)

    # Comment handling
    def _read_comments_xml(self, docx_path: str) -> Dict[str, Dict[str, Any]]:
        """Return {comment_id: {id, author, date, text}}."""
        try:
            with zipfile.ZipFile(docx_path) as z:
                try:
                    xml_bytes = z.read("word/comments.xml")
                except KeyError:
                    return {}
            root = ET.fromstring(xml_bytes)

            comments: Dict[str, Dict[str, Any]] = {}
            for c in root.findall("w:comment", self.NS):
                cid = c.get(f"{{{self.W_NS}}}id")
                author = c.get(f"{{{self.W_NS}}}author")
                date = c.get(f"{{{self.W_NS}}}date")
                texts = [t.text for t in c.findall(".//w:t", self.NS) if t.text]
                comment_text = "".join(texts).strip()
                if cid is not None:
                    comments[cid] = {"id": cid, "author": author, "date": date, "text": comment_text}
            return comments
        except Exception:
            return {}

    # Comment anchor handling
    def _extract_comments_with_anchors(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        Extract comments + where they attach:
          - anchor_heading_path
          - anchor_text (text within comment range)
          - anchor_context (paragraph-level context)
        """
        comment_meta = self._read_comments_xml(docx_path)
        if not comment_meta:
            return []

        doc = Document(docx_path)

        current_path: List[str] = []
        acc: Dict[str, Dict[str, Any]] = {
            cid: {
                "anchor_heading_path": [],
                "anchor_text_parts": [],
                "anchor_text_rich_parts": [],
                "anchor_context_parts": [],
                "anchor_context_rich_parts": [],
            }
            for cid in comment_meta.keys()
        }

        active_across: set[str] = set()

        def update_heading(p: Paragraph) -> bool:
            nonlocal current_path
            lvl = self._heading_level(p)
            if lvl not in (1, 2):
                return False
            title = (p.text or "").strip()
            if not title:
                return True
            if lvl == 1:
                current_path = [title]
            else:
                current_path = (current_path[:1] if current_path else []) + [title]
            return True

        def scan_paragraph(p: Paragraph) -> None:
            nonlocal active_across

            if update_heading(p):
                return

            paragraph_touched: set[str] = set()
            active_stack: List[str] = list(active_across)

            def ensure_heading(cid: str) -> None:
                if not acc[cid]["anchor_heading_path"]:
                    acc[cid]["anchor_heading_path"] = list(current_path)

            for child in p._p.iterchildren():
                tag = child.tag

                if tag == f"{{{self.W_NS}}}commentRangeStart":
                    cid = child.get(f"{{{self.W_NS}}}id")
                    if cid in acc:
                        active_stack.append(cid)
                        active_across.add(cid)
                        ensure_heading(cid)
                        paragraph_touched.add(cid)

                elif tag == f"{{{self.W_NS}}}commentRangeEnd":
                    cid = child.get(f"{{{self.W_NS}}}id")
                    if cid in acc:
                        for i in range(len(active_stack) - 1, -1, -1):
                            if active_stack[i] == cid:
                                active_stack.pop(i)
                                break
                        if cid not in active_stack:
                            active_across.discard(cid)
                        paragraph_touched.add(cid)

                elif tag == f"{{{self.W_NS}}}r":
                    texts = [t.text for t in child.findall(".//w:t", self.NS) if t.text]
                    chunk = "".join(texts)
                    chunk_rich = self._xml_run_text_with_scripts(child)
                    if chunk:
                        for cid in active_stack:
                            if cid in acc:
                                ensure_heading(cid)
                                acc[cid]["anchor_text_parts"].append(chunk)
                                acc[cid]["anchor_text_rich_parts"].append(chunk_rich or chunk)
                                paragraph_touched.add(cid)

            # Keep a short paragraph-level context for each touched comment.
            ptxt = (p.text or "").strip()
            if ptxt:
                for cid in (paragraph_touched | set(active_stack)):
                    if cid in acc and len(acc[cid]["anchor_context_parts"]) < 6:
                        acc[cid]["anchor_context_parts"].append(ptxt)

            ptxt_rich = self._paragraph_text_with_scripts(p)
            if ptxt_rich:
                for cid in (paragraph_touched | set(active_stack)):
                    if cid in acc:
                        if len(acc[cid]["anchor_context_rich_parts"]) < 6:
                            acc[cid]["anchor_context_rich_parts"].append(ptxt_rich)

        for item in self._iter_block_items_in_order(doc):
            if isinstance(item, Paragraph):
                scan_paragraph(item)
            elif isinstance(item, Table):
                for row in item.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            scan_paragraph(p)

        out: List[Dict[str, Any]] = []
        for cid, meta in comment_meta.items():
            out.append(
                {
                    **meta,
                    "anchor_heading_path": acc[cid]["anchor_heading_path"],
                    "anchor_text": "".join(acc[cid]["anchor_text_parts"]).strip(),
                    "anchor_text_rich": "".join(acc[cid]["anchor_text_rich_parts"]).strip(),
                    "anchor_context": " ".join(acc[cid]["anchor_context_parts"]).strip(),
                    "anchor_context_rich": " ".join(acc[cid]["anchor_context_rich_parts"]).strip(),
                }
            )
        return out
    
    def _parse_html(self, html_path: str) -> Dict[str, Any]:
        """
        Parse an HTML file into the same dict schema as DOCX.

        Comments: HTML typically has none => [].
        """
        html = Path(html_path).read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")

        # Discover class names that imply bold/italic from embedded CSS.
        css_text = ""
        style_tag = soup.find("style")
        if style_tag:
            css_text = style_tag.get_text(" ", strip=True).lower()

        bold_classes = set()
        italic_classes = set()

        # Very lightweight CSS parsing: .class { ... font-weight: bold ... }
        import re
        for m in re.finditer(r"\.(?P<cls>[a-z0-9_-]+)\s*\{[^}]*font-weight\s*:\s*bold", css_text):
            bold_classes.add(m.group("cls"))

        for m in re.finditer(r"\.(?P<cls>[a-z0-9_-]+)\s*\{[^}]*font-style\s*:\s*italic", css_text):
            italic_classes.add(m.group("cls"))


        doc_title = Path(html_path).stem
        root = HeadingNode(title=doc_title, level=0, path=[], blocks=[], children=[])

        current_h1: Optional[HeadingNode] = None
        current_h2: Optional[HeadingNode] = None

        def container() -> HeadingNode:
            return current_h2 or current_h1 or root

        def style_bucket_for(node: HeadingNode) -> Dict[str, List[str]]:
            if not hasattr(node, "_style_bucket"):
                setattr(node, "_style_bucket", self._empty_style_bucket())
            return getattr(node, "_style_bucket")

        def set_heading(title: str, lvl: int) -> None:
            nonlocal current_h1, current_h2
            title = (title or "").strip()
            if not title:
                return

            if lvl == 1:
                node = HeadingNode(title=title, level=1, path=[title], blocks=[], children=[])
                root.children.append(node)
                current_h1, current_h2 = node, None
                return

            if lvl == 2:
                if current_h1 is None:
                    node = HeadingNode(title=title, level=2, path=[title], blocks=[], children=[])
                    root.children.append(node)
                    current_h2 = node
                else:
                    node = HeadingNode(title=title, level=2, path=current_h1.path + [title], blocks=[], children=[])
                    current_h1.children.append(node)
                    current_h2 = node

        def clean_text(el) -> str:
            return " ".join(el.get_text(" ", strip=True).split())

        # Iterate in DOM order: headings, paragraphs, lists, and tables.
        for el in soup.find_all(["h1", "h2", "p", "ul", "ol", "table"]):
            name = el.name.lower()

            if name == "h1":
                set_heading(clean_text(el), 1)
                continue
            if name == "h2":
                set_heading(clean_text(el), 2)
                continue

            if name == "p":
                text = clean_text(el)
                if text:
                    self._merge_style_bucket(style_bucket_for(container()), self._extract_styles_from_html_element(el, bold_classes=bold_classes, italic_classes=italic_classes))
                    container().blocks.append({"type": "paragraph", "text": text})
                continue

            if name in ("ul", "ol"):
                items = []
                for li in el.find_all("li", recursive=False):
                    t = clean_text(li)
                    self._merge_style_bucket(style_bucket_for(container()), self._extract_styles_from_html_element(li))
                    if t:
                        items.append({"text": t, "style": f"HTML:{name}/li"})
                if items:
                    container().blocks.append({"type": "list", "signature": f"HTML:{name}", "items": items})
                continue

            if name == "table":
                self._merge_style_bucket(style_bucket_for(container()), self._extract_styles_from_html_element(el, bold_classes=bold_classes, italic_classes=italic_classes))
                rows: List[List[str]] = []
                for tr in el.find_all("tr"):
                    row = []
                    for cell in tr.find_all(["th", "td"]):
                        row.append(clean_text(cell))
                    if row:
                        rows.append(row)
                if rows:
                    container().blocks.append({"type": "table", "rows": rows})
                continue

        def add_style_block_if_any(node: HeadingNode) -> None:
            bucket = getattr(node, "_style_bucket", None)
            if bucket and any(bucket[k] for k in bucket):
                node.blocks.append({"type": "style", "data": bucket})
            for ch in node.children:
                add_style_block_if_any(ch)

        add_style_block_if_any(root)

        out = root.to_dict()
        out["comments"] = []  # no DOCX comment anchors in plain HTML
        return out



# Batch processing
def run_batch_default(parser: AssessmentParser) -> int:
    """Batch mode using DEFAULT_INPUT_FOLDER and DEFAULT_OUTPUT_FOLDER."""
    in_path = Path(DEFAULT_INPUT_FOLDER)
    out_path = Path(DEFAULT_OUTPUT_FOLDER)

    if not in_path.exists() or not in_path.is_dir():
        print(f"Input folder not found: {in_path}")
        return 1

    out_path.mkdir(parents=True, exist_ok=True)

    files = sorted(
    p for p in in_path.iterdir()
    if p.is_file() and p.suffix.lower() in (".docx", ".html", ".htm")
    )
    if not files:
        print(f"No .docx files found in {in_path}")
        return 0

    errors: List[Dict[str, str]] = []
    parsed = 0

    for file in files:
        try:
            data = parser.parse_file(str(file))
            json_file = out_path / f"{file.stem}.json"
            with open(json_file, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            parsed += 1
            print(f"Done: {file.name}")
        except Exception as e:
            errors.append({"file": file.name, "error": repr(e)})
            print(f"Failed: {file.name} ({e!r})")

    errors_file = out_path / "_errors.json"
    with open(errors_file, "w", encoding="utf-8") as f:
        json.dump(
            {
                "input_folder": str(in_path),
                "output_folder": str(out_path),
                "total": len(files),
                "parsed": parsed,
                "failed": len(errors),
                "errors": errors,
            },
            f,
            indent=2,
            ensure_ascii=False,
        )

    print(f"\nDone: {parsed}/{len(files)} parsed. Errors: {errors_file.name}")
    return 0


def parse_dict(docx_path: str) -> Dict[str, Any]:
    parser = AssessmentParser()
    return parser.parse_file(docx_path)


def parse_to_dict(docx_path: str) -> Dict[str, Any]:
    return parse_dict(docx_path)


def main() -> int:
    """
    - No args: run batch with default folders
    - One arg ending with .docx: parse that file and print JSON to stdout
    """
    parser = AssessmentParser()
    args = sys.argv[1:]

    if not args:
        return run_batch_default(parser)

    if len(args) == 1 and args[0].lower().endswith((".docx", ".html", ".htm")):
        data = parser.parse_file(args[0])
        print(json.dumps(data, indent=2, ensure_ascii=False))
        return 0

    return 1


if __name__ == "__main__":
    raise SystemExit(main())
