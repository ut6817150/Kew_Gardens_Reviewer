"""
assessment_processor.py

Converts a word document of an assessment draft to a json dict.

Produces (per file):
- Heading skeleton (Heading 1/Heading 2) in document order
- Under each heading: ordered blocks[] of typed items: paragraph, list, table
- Dcomment comments: including context 

Run:
  (A) Default batch mode:
      python3.12 assessment_processor.py

  (B) Single-file debug mode:
      python3.12 assessment_processor.py "My Assessment.docx"
"""

# Imports: 
from __future__ import annotations
import json
import sys
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple, Union
from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph


# Default folders in the same directory
DEFAULT_INPUT_FOLDER = "converted"
DEFAULT_OUTPUT_FOLDER = "json_converted"


# Create the tree structure:
@dataclass
class HeadingNode:
    """A heading node in the skeleton tree."""
    title: Optional[str]
    level: int
    path: List[str]
    blocks: List[Dict[str, Any]]
    children: List["HeadingNode"]

    def to_dict(self) -> Dict[str, Any]:
        return {
            "title": self.title,
            "level": self.level,
            "path": list(self.path),
            "blocks": list(self.blocks),
            "children": [c.to_dict() for c in self.children],
        }


class AssessmentParser:
    """
    Document parser class.

    The class goes through a word file to:
      - Build heading skeleton 
      - Attach blocks to headings (paragraph/list/table)
      - Extract comments 
    """

    ## WordprocessingML namespace used by Microsoft Word (.docx).
    # Required to correctly find comments and comment anchors in the XML;
    # without this, ElementTree searches would return nothing.
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS = {"w": W_NS}

    def parse_file(self, docx_path: str) -> Dict[str, Any]:
        """Main file parser, takes one DOCX file and produces a JSON dict."""
        doc = Document(docx_path)

        root = self._build_heading_skeleton(doc)
        self._attach_blocks(doc, root)

        out = root.to_dict()
        out["comments"] = self._extract_comments_with_anchors(docx_path)
        return out


    # DOCX traversal:
    def _iter_block_items_in_order(self, doc: _Document) -> Iterable[Union[Paragraph, Table]]:
        """Yield top-level paragraphs and tables in the order they appear."""
        for child in doc.element.body.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, doc)
            elif child.tag.endswith("}tbl"):
                yield Table(child, doc)

    def _paragraph_style_name(self, p: Paragraph) -> str:
        try:
            return (p.style.name or "").strip()
        except Exception:
            return ""

    def _heading_level(self, p: Paragraph) -> Optional[int]:
        """Return 1/2 if style is Heading 1/Heading 2 else None."""
        name = self._paragraph_style_name(p)
        if name == "Heading 1" or name.replace(" ", "").lower() == "heading1":
            return 1
        if name == "Heading 2" or name.replace(" ", "").lower() == "heading2":
            return 2
        return None

    def _is_list_item(self, p: Paragraph) -> bool:
        """Detect Word list items via <w:numPr>."""
        ppr = p._p.pPr
        return ppr is not None and ppr.numPr is not None

    def _list_signature(self, p: Paragraph) -> str:
        """Group consecutive list items by numId/ilvl where possible."""
        ppr = p._p.pPr
        style = self._paragraph_style_name(p)

        if ppr is not None and ppr.numPr is not None:
            numPr = ppr.numPr
            numId = numPr.numId.val if numPr.numId is not None else None
            ilvl = numPr.ilvl.val if numPr.ilvl is not None else None
            return f"numId={numId}|ilvl={ilvl}|style={style}"

        return f"style={style}"

    def _table_to_rows(self, tbl: Table) -> List[List[str]]:
        """Extract table into 2D list of cell text."""
        rows: List[List[str]] = []
        for row in tbl.rows:
            out_row: List[str] = []
            for cell in row.cells:
                text = "\n".join(
                    [p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()]
                )
                out_row.append(text)
            rows.append(out_row)
        return rows


    # Header builder:
    def _build_heading_skeleton(self, doc: _Document) -> HeadingNode:
        """
        Build Heading 1/Heading 2 skeleton in document order.
        """
        root = HeadingNode(title=None, level=0, path=[], blocks=[], children=[])
        current_h1: Optional[HeadingNode] = None

        def add_heading(title: str, lvl: int) -> None:
            nonlocal current_h1
            title = title.strip()
            if not title:
                return

            if lvl == 1:
                node = HeadingNode(title=title, level=1, path=[title], blocks=[], children=[])
                root.children.append(node)
                current_h1 = node
                return

            if lvl == 2:
                parent = current_h1
                node_path = [title] if parent is None else parent.path + [title]
                node = HeadingNode(title=title, level=2, path=node_path, blocks=[], children=[])
                if parent is None:
                    root.children.append(node)  # orphan H2
                else:
                    parent.children.append(node)

        for item in self._iter_block_items_in_order(doc):
            if isinstance(item, Paragraph):
                lvl = self._heading_level(item)
                if lvl in (1, 2):
                    add_heading(item.text, lvl)

        return root

    # Blocks attacher:
    def _attach_blocks(self, doc: _Document, root: HeadingNode) -> None:
        """
        Walk the doc in order and attach paragraph/list/table blocks to the
        current heading (deepest of H2 then H1; else root).
        """
        current_h1: Optional[HeadingNode] = None
        current_h2: Optional[HeadingNode] = None

        def container() -> HeadingNode:
            if current_h2 is not None:
                return current_h2
            if current_h1 is not None:
                return current_h1
            return root

        def set_heading(title: str, lvl: int) -> None:
            nonlocal current_h1, current_h2
            title = title.strip()
            if not title:
                return

            if lvl == 1:
                for n in root.children:
                    if n.level == 1 and n.title == title:
                        current_h1, current_h2 = n, None
                        return
                # fallback create
                n = HeadingNode(title=title, level=1, path=[title], blocks=[], children=[])
                root.children.append(n)
                current_h1, current_h2 = n, None
                return

            if lvl == 2:
                parent = current_h1
                if parent is None:
                    # orphan H2
                    for n in root.children:
                        if n.level == 2 and n.title == title:
                            current_h2 = n
                            return
                    n = HeadingNode(title=title, level=2, path=[title], blocks=[], children=[])
                    root.children.append(n)
                    current_h2 = n
                    return

                for n in parent.children:
                    if n.level == 2 and n.title == title:
                        current_h2 = n
                        return
                n = HeadingNode(title=title, level=2, path=parent.path + [title], blocks=[], children=[])
                parent.children.append(n)
                current_h2 = n

        pending_list: Optional[Dict[str, Any]] = None
        pending_sig: Optional[str] = None

        def flush_list() -> None:
            nonlocal pending_list, pending_sig
            if pending_list is not None:
                container().blocks.append(pending_list)
                pending_list = None
                pending_sig = None

        for item in self._iter_block_items_in_order(doc):
            if isinstance(item, Paragraph):
                lvl = self._heading_level(item)
                if lvl in (1, 2):
                    flush_list()
                    set_heading(item.text, lvl)
                    continue

                text = (item.text or "").strip()
                if not text:
                    continue

                if self._is_list_item(item):
                    sig = self._list_signature(item)
                    if pending_list is None or sig != pending_sig:
                        flush_list()
                        pending_list = {"type": "list", "signature": sig, "items": []}
                        pending_sig = sig
                    pending_list["items"].append({"text": text, "style": self._paragraph_style_name(item)})
                else:
                    flush_list()
                    container().blocks.append({"type": "paragraph", "text": text, "style": self._paragraph_style_name(item)})

            elif isinstance(item, Table):
                flush_list()
                container().blocks.append({"type": "table", "rows": self._table_to_rows(item)})

        flush_list()


    # Comment handling:
    def _read_comments_xml(self, docx_path: str) -> Dict[str, Dict[str, Any]]:
        """Return {comment_id: {id, author, date, text}}."""
        try:
            with zipfile.ZipFile(docx_path) as z:
                try:
                    xml_bytes = z.read("word/comments.xml")
                except KeyError:
                    return {}
            root = ET.fromstring(xml_bytes)

            comments: Dict[str, Dict[str, Any]] = {}
            for c in root.findall("w:comment", self.NS):
                cid = c.get(f"{{{self.W_NS}}}id")
                author = c.get(f"{{{self.W_NS}}}author")
                date = c.get(f"{{{self.W_NS}}}date")
                texts = [t.text for t in c.findall(".//w:t", self.NS) if t.text]
                comment_text = "".join(texts).strip()
                if cid is not None:
                    comments[cid] = {"id": cid, "author": author, "date": date, "text": comment_text}
            return comments
        except Exception:
            return {}

    # comment anchers handling:
    def _extract_comments_with_anchors(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        Extract comments + where they attach:
          - anchor_heading_path
          - anchor_text (text within comment range)
          - anchor_context (paragraph-level context)
        """
        comment_meta = self._read_comments_xml(docx_path)
        if not comment_meta:
            return []

        doc = Document(docx_path)

        current_path: List[str] = []
        acc: Dict[str, Dict[str, Any]] = {
            cid: {"anchor_heading_path": [], "anchor_text_parts": [], "anchor_context_parts": []}
            for cid in comment_meta.keys()
        }

        active_across: set[str] = set()

        def update_heading(p: Paragraph) -> bool:
            nonlocal current_path
            lvl = self._heading_level(p)
            if lvl not in (1, 2):
                return False
            title = (p.text or "").strip()
            if not title:
                return True
            if lvl == 1:
                current_path = [title]
            else:
                current_path = (current_path[:1] if current_path else []) + [title]
            return True

        def scan_paragraph(p: Paragraph) -> None:
            nonlocal active_across

            if update_heading(p):
                return

            paragraph_touched: set[str] = set()
            active_stack: List[str] = list(active_across)

            def ensure_heading(cid: str) -> None:
                if not acc[cid]["anchor_heading_path"]:
                    acc[cid]["anchor_heading_path"] = list(current_path)

            for child in p._p.iterchildren():
                tag = child.tag

                if tag == f"{{{self.W_NS}}}commentRangeStart":
                    cid = child.get(f"{{{self.W_NS}}}id")
                    if cid in acc:
                        active_stack.append(cid)
                        active_across.add(cid)
                        ensure_heading(cid)
                        paragraph_touched.add(cid)

                elif tag == f"{{{self.W_NS}}}commentRangeEnd":
                    cid = child.get(f"{{{self.W_NS}}}id")
                    if cid in acc:
                        for i in range(len(active_stack) - 1, -1, -1):
                            if active_stack[i] == cid:
                                active_stack.pop(i)
                                break
                        if cid not in active_stack:
                            active_across.discard(cid)
                        paragraph_touched.add(cid)

                elif tag == f"{{{self.W_NS}}}r":
                    texts = [t.text for t in child.findall(".//w:t", self.NS) if t.text]
                    chunk = "".join(texts)
                    if chunk:
                        for cid in active_stack:
                            if cid in acc:
                                ensure_heading(cid)
                                acc[cid]["anchor_text_parts"].append(chunk)
                                paragraph_touched.add(cid)

            # paragraph-level context
            ptxt = (p.text or "").strip()
            if ptxt:
                for cid in (paragraph_touched | set(active_stack)):
                    if cid in acc and len(acc[cid]["anchor_context_parts"]) < 6:
                        acc[cid]["anchor_context_parts"].append(ptxt)

        for item in self._iter_block_items_in_order(doc):
            if isinstance(item, Paragraph):
                scan_paragraph(item)
            elif isinstance(item, Table):
                for row in item.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            scan_paragraph(p)

        out: List[Dict[str, Any]] = []
        for cid, meta in comment_meta.items():
            out.append(
                {
                    **meta,
                    "anchor_heading_path": acc[cid]["anchor_heading_path"],
                    "anchor_text": "".join(acc[cid]["anchor_text_parts"]).strip(),
                    "anchor_context": " ".join(acc[cid]["anchor_context_parts"]).strip(),
                }
            )
        return out


# Processor for multible documents:
def run_batch_default(parser: AssessmentParser) -> int:
    """Batch mode using DEFAULT_INPUT_FOLDER and DEFAULT_OUTPUT_FOLDER."""
    in_path = Path(DEFAULT_INPUT_FOLDER)
    out_path = Path(DEFAULT_OUTPUT_FOLDER)

    if not in_path.exists() or not in_path.is_dir():
        print(f"Input folder not found: {in_path}")
        return 1

    out_path.mkdir(parents=True, exist_ok=True)

    docx_files = sorted(p for p in in_path.iterdir() if p.is_file() and p.suffix.lower() == ".docx")
    if not docx_files:
        print(f"No .docx files found in {in_path}")
        return 0

    errors: List[Dict[str, str]] = []
    parsed = 0

    for docx_file in docx_files:
        try:
            data = parser.parse_file(str(docx_file))
            json_file = out_path / f"{docx_file.stem}.json"
            with open(json_file, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            parsed += 1
            print(f"Done: {docx_file.name}")
        except Exception as e:
            errors.append({"file": docx_file.name, "error": repr(e)})
            print(f"Failed: {docx_file.name} ({e!r})")

    errors_file = out_path / "_errors.json"
    with open(errors_file, "w", encoding="utf-8") as f:
        json.dump(
            {
                "input_folder": str(in_path),
                "output_folder": str(out_path),
                "total": len(docx_files),
                "parsed": parsed,
                "failed": len(errors),
                "errors": errors,
            },
            f,
            indent=2,
            ensure_ascii=False,
        )

    print(f"\nDone: {parsed}/{len(docx_files)} parsed. Errors: {errors_file.name}")
    return 0


def parse_docx_to_dict(docx_path: str) -> Dict[str, Any]:
    parser = AssessmentParser()
    return parser.parse_file(docx_path)


def main() -> int:
    """
    - No args: run batch with default folders
    - One arg ending with .docx: parse that file and print JSON to stdout
    """
    parser = AssessmentParser()
    args = sys.argv[1:]

    if not args:
        return run_batch_default(parser)

    if len(args) == 1 and args[0].lower().endswith(".docx"):
        data = parser.parse_file(args[0])
        print(json.dumps(data, indent=2, ensure_ascii=False))
        return 0

    return 1


if __name__ == "__main__":
    raise SystemExit(main())
